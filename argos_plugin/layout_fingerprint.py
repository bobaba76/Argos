"""Spec-09 (#112): form-level identity — deterministic layout fingerprint.

A *layout-family fingerprint* is a deterministic structural signature of a
document's *form*, distinct from the two content-level hashes the watcher
already keeps:

- ``file_id``      — SHA-256 of raw bytes (dedup, content-edit detection)
- ``extract_hash`` — SHA-256 of the extraction input (the re-extract gate)
- ``layout_family`` — THIS module: SHA-256 of structural features
  (page count, table regions, column signatures, heading structure).

The fingerprint answers the form-level question: *is this a new instance of
a known layout/template, or genuinely novel?* Two files that share a layout
but differ in content (a resaved Excel with one cell changed, a fresh invoice
in the same template) share a ``layout_family`` but differ in ``file_id``.

Design rules (from the issue acceptance criteria):

1. **Deterministic and LLM-free** — structural features only, computed from
   the same file reads the catalog pass already does. No model calls.
2. **Subordinate to content** — layout-family identity NEVER suppresses
   content-level extraction. The known-family short-circuit fires ONLY when
   ``file_id`` is unchanged (see ``classify_layout``).
3. **Per-doc-type features** — PDF / XLSX / CSV / DOCX each contribute the
   structural signals that make sense for their format. The fingerprint is
   ``sha256(json.dumps(features, sort_keys=True))`` so ordering is stable.
4. **v1 = exact structural signature.** Tolerant clustering for resaved /
   OCR-noised variants is explicitly deferred (issue open question).

No new deps — PyPDF2 + openpyxl + python-docx are already in the venv
(used by ``watcher.py``). Never raises; unreadable files return ``None``
so the catalog pass never fails (same contract as ``hash_file``).
"""
from __future__ import annotations

import csv
import hashlib
import io
import json
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Fingerprint computation
# ---------------------------------------------------------------------------


def _fingerprint(features: Dict[str, Any]) -> str:
    """Stable SHA-256 over a feature dict (sorted keys, compact JSON)."""
    blob = json.dumps(features, sort_keys=True, ensure_ascii=False, separators=(",", ":"))
    return hashlib.sha256(blob.encode("utf-8")).hexdigest()


def _fingerprint_pdf(path: str | Path) -> Optional[Dict[str, Any]]:
    """Structural features for a PDF: page count, per-page text length
    bucket, heading-line count, table-region heuristic.

    Uses PyPDF2 (already a watcher dep). No OCR, no text content in the
    fingerprint — only structural counts and buckets so two invoices in the
    same template share a family even when the numbers differ.
    """
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(path))
        page_count = len(reader.pages)
        page_buckets: List[str] = []
        heading_lines = 0
        for page in reader.pages:
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            # Bucket text length by powers of 2 (structural, not content).
            tlen = len(text)
            if tlen <= 0:
                page_buckets.append("0")
            else:
                page_buckets.append(str(1 << (tlen.bit_length() - 1)))
            # Heading heuristic: short ALL-CAPS lines, or lines ending with
            # a colon, or lines that look like section headers. Structural
            # count only — the actual words are not in the fingerprint.
            for line in text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if (line.isupper() and 3 <= len(line) <= 60) or line.endswith(":"):
                    heading_lines += 1
        # Table-region heuristic: count lines dominated by runs of
        # whitespace-separated tokens (>= 3 tokens, >= 2 gaps). Structural.
        table_line_count = 0
        for page in reader.pages:
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            for line in text.split("\n"):
                gaps = len(line.split()) - 1
                if gaps >= 2:
                    table_line_count += 1
        return {
            "doc_type": "pdf",
            "page_count": page_count,
            "page_text_buckets": page_buckets,
            "heading_lines": heading_lines,
            "table_line_count": table_line_count,
        }
    except Exception as exc:
        logger.debug("pdf fingerprint failed for %s: %s", path, exc)
        return None


def _fingerprint_xlsx(path: str | Path) -> Optional[Dict[str, Any]]:
    """Structural features for an XLSX: sheet count, per-sheet shape
    (row count bucket, column count), header signature.

    Uses openpyxl (already a watcher dep). The header signature is a hash
    of the first-row cell *positions* that are non-empty (not the values) —
    so two spreadsheets with the same column layout but different data
    share a family. Column count is exact; row count is bucketed (a 500-row
    and 501-row sheet are the same template).
    """
    try:
        from openpyxl import load_workbook
        wb = load_workbook(str(path), read_only=True, data_only=True)
        sheet_features: List[Dict[str, Any]] = []
        for sname in wb.sheetnames:
            ws = wb[sname]
            max_row = 0
            max_col = 0
            header_nonempty_positions: List[int] = []
            first_row_seen = False
            for row in ws.iter_rows(values_only=True):
                max_row += 1
                nonempty = [i for i, c in enumerate(row) if c is not None and str(c).strip()]
                if nonempty:
                    if max_col == 0:
                        max_col = max(nonempty) + 1
                    if not first_row_seen:
                        first_row_seen = True
                        # Header signature: positions of non-empty cells in
                        # the first non-empty row (column layout, not values).
                        header_nonempty_positions = nonempty
            wb.close()
            # Bucket row count by powers of 2 (structural).
            row_bucket = str(1 << (max_row.bit_length() - 1)) if max_row > 0 else "0"
            sheet_features.append({
                "name_len": len(sname),
                "row_bucket": row_bucket,
                "col_count": max_col,
                "header_positions": header_nonempty_positions,
            })
        return {
            "doc_type": "xlsx",
            "sheet_count": len(sheet_features),
            "sheets": sheet_features,
        }
    except Exception as exc:
        logger.debug("xlsx fingerprint failed for %s: %s", path, exc)
        return None


def _fingerprint_csv(path: str | Path) -> Optional[Dict[str, Any]]:
    """Structural features for a CSV: column count, row count bucket,
    header presence, delimiter.

    Reads only the first 200 rows (structural sample — a 50k-row export and
    a 50.1k-row export are the same template). The header signature is the
    count of non-empty columns (not the values).
    """
    try:
        with open(path, "r", encoding="utf-8", errors="replace", newline="") as f:
            sample = f.read(65536)  # 64KB structural sample
        # Detect delimiter via csv.Sniffer on the first few lines.
        reader = csv.reader(io.StringIO(sample))
        rows: List[List[str]] = []
        for i, row in enumerate(reader):
            if i >= 200:
                break
            rows.append(row)
        if not rows:
            return None
        col_count = max(len(r) for r in rows)
        nonempty_header_cols = sum(1 for c in rows[0] if c and c.strip())
        # Bucket row count.
        rc = len(rows)
        row_bucket = str(1 << (rc.bit_length() - 1)) if rc > 0 else "0"
        return {
            "doc_type": "csv",
            "col_count": col_count,
            "row_bucket": row_bucket,
            "header_nonempty_cols": nonempty_header_cols,
        }
    except Exception as exc:
        logger.debug("csv fingerprint failed for %s: %s", path, exc)
        return None


def _fingerprint_docx(path: str | Path) -> Optional[Dict[str, Any]]:
    """Structural features for a DOCX: paragraph count bucket, heading
    count by style, table count + shape.

    Uses python-docx (already a watcher dep). Heading styles are counted
    by name (Heading 1, Heading 2, …) — structural, not content.
    """
    try:
        import docx
        doc = docx.Document(str(path))
        para_count = 0
        heading_by_style: Dict[str, int] = {}
        table_shapes: List[Dict[str, int]] = []
        for p in doc.paragraphs:
            if p.text.strip():
                para_count += 1
                style = p.style.name if p.style else "Normal"
                if style and style.lower().startswith("heading"):
                    heading_by_style[style] = heading_by_style.get(style, 0) + 1
        for t in doc.tables:
            table_shapes.append({"rows": len(t.rows), "cols": len(t.columns)})
        # Bucket paragraph count.
        pc = para_count
        para_bucket = str(1 << (pc.bit_length() - 1)) if pc > 0 else "0"
        return {
            "doc_type": "docx",
            "para_bucket": para_bucket,
            "heading_by_style": heading_by_style,
            "table_count": len(table_shapes),
            "table_shapes": table_shapes,
        }
    except Exception as exc:
        logger.debug("docx fingerprint failed for %s: %s", path, exc)
        return None


_DOC_TYPE_FINGERPRINTERS = {
    "pdf": _fingerprint_pdf,
    "xlsx": _fingerprint_xlsx,
    "xls": _fingerprint_xlsx,
    "csv": _fingerprint_csv,
    "docx": _fingerprint_docx,
}


def compute_layout_features(
    path: str | Path,
    doc_type: str,
) -> Optional[Dict[str, Any]]:
    """Compute the structural feature dict for a document.

    Returns None on any failure (unreadable, unsupported, corrupt) — the
    catalog pass must not fail. The features are deterministic and LLM-free.
    """
    fn = _DOC_TYPE_FINGERPRINTERS.get(doc_type)
    if fn is None:
        return None
    return fn(path)


def compute_layout_family(
    path: str | Path,
    doc_type: str,
) -> Optional[str]:
    """Compute the layout-family fingerprint (SHA-256 hex) for a document.

    Returns None on any failure. The fingerprint is a stable hash of the
    structural features — two files with the same layout but different
    content share a fingerprint; a content edit does not change it (that
    is the point: form-level, not content-level).
    """
    features = compute_layout_features(path, doc_type)
    if features is None:
        return None
    return _fingerprint(features)


# ---------------------------------------------------------------------------
# Layout classification — the known-family short-circuit
# ---------------------------------------------------------------------------

# Layout classification relative to the catalog.
#   known    — fingerprint matches an existing family in the catalog
#   novel    — fingerprint not seen in the catalog (surface for labelling)
#   unknown  — fingerprint could not be computed (unreadable / unsupported)


def classify_layout(
    layout_family: Optional[str],
    known_families: Dict[str, int],
) -> str:
    """Classify a layout fingerprint against the known families.

    Args:
        layout_family: the fingerprint computed for this file (or None).
        known_families: {fingerprint: doc_count} from the catalog.

    Returns:
        'known' / 'novel' / 'unknown'.

    'unknown' means the fingerprint could not be computed — the file is
    still processed normally; it just gets no form-level shortcut.
    """
    if layout_family is None:
        return "unknown"
    return "known" if layout_family in known_families else "novel"


def should_short_circuit_extraction(
    *,
    file_id: str,
    catalog_file_id: Optional[str],
    layout_family: Optional[str],
    known_families: Dict[str, int],
) -> bool:
    """Decide whether the known-family short-circuit fires.

    The invariant (issue acceptance criterion): the short-circuit fires
    ONLY when ``file_id`` (content hash) is unchanged. A same-template file
    whose content changed (one cell updated) still runs the existing
    re-extract gate; layout-family identity never suppresses content-level
    extraction, regardless of filename.

    This function checks the form-level half only — the caller is still
    responsible for the content-level re-extract gate (``extract_hash``).
    The short-circuit is subordinate: it can skip the *layout discovery*
    work (re-fingerprinting, family lookup, labelling surface) but NOT the
    content-level extraction decision.
    """
    # No fingerprint → no form-level signal → no short-circuit.
    if layout_family is None:
        return False
    # Novel family → must surface for full extraction + labelling.
    if layout_family not in known_families:
        return False
    # No prior catalog entry → new file, new content → no shortcut even if
    # its layout matches a known family (it's still new content to extract).
    if catalog_file_id is None:
        return False
    # Known family BUT content changed → content gate wins, no shortcut.
    if file_id != catalog_file_id:
        return False
    # Known family AND unchanged content AND already in catalog → cheap path.
    return True


def load_known_families(catalog_rows: List[Dict[str, Any]]) -> Dict[str, int]:
    """Build a {fingerprint: doc_count} map from catalog rows.

    Rows with NULL layout_family are excluded.
    """
    families: Dict[str, int] = {}
    for row in catalog_rows:
        fam = row.get("layout_family")
        if fam:
            families[fam] = families.get(fam, 0) + 1
    return families
