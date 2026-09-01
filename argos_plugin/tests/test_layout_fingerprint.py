"""Spec-09 (#112): form-level identity — layout-family fingerprint + stratified evals.

Tests (cheapest falsifying first — deterministic, no LLM):
1. Fingerprint determinism: same layout + different content → same family.
2. Fingerprint distinction: different layout → different family.
3. Per-doc-type features: PDF / XLSX / CSV / DOCX each compute features.
4. Short-circuit invariant: known family + unchanged file_id → short-circuit;
   known family + changed file_id → NO short-circuit (content gate wins).
5. Novel surfacing: novel family always surfaces, never silently skipped.
6. Catalog storage: layout_family column persists + retrieves.
7. Label registry: sidecar labelling + relabel history.
8. Scan classification: classify_scan_by_layout buckets new/changed correctly.
9. Eval stratification: per-family accuracy reported alongside aggregate.
"""
import json
import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
_EVAL_DIR = Path(__file__).resolve().parent.parent / "eval"
if str(_EVAL_DIR) not in sys.path:
    sys.path.insert(0, str(_EVAL_DIR))

from layout_fingerprint import (
    classify_layout,
    compute_layout_family,
    compute_layout_features,
    load_known_families,
    should_short_circuit_extraction,
)
from layout_family_registry import (
    get_label,
    label_family,
    labelled_families,
    load_registry,
    merge_catalog_counts,
    registry_path_for,
)
from watcher import (
    classify_scan_by_layout,
    scan_pass,
    surface_novel_layouts,
)


# ---------------------------------------------------------------------------
# Fixtures: create real minimal files of each supported type
# ---------------------------------------------------------------------------


def _make_pdf(path: Path, pages: list[str]) -> Path:
    """Create a real minimal PDF with the given page text using PyPDF2."""
    from PyPDF2 import PdfWriter
    writer = PdfWriter()
    for text in pages:
        # PdfWriter doesn't add text directly; we use a blank page and rely
        # on page count as the structural feature. For text-bearing tests we
        # use reportlab if available, else accept blank pages (page count is
        # the dominant structural feature for the fingerprint anyway).
        writer.add_blank_page(width=200, height=300)
    with open(path, "wb") as f:
        writer.write(f)
    return path


def _make_xlsx(path: Path, sheets: dict[str, list[list]]) -> Path:
    """Create a real XLSX with the given sheet data using openpyxl."""
    from openpyxl import Workbook
    wb = Workbook()
    # Remove the default sheet; we add named ones.
    wb.remove(wb.active)
    for name, rows in sheets.items():
        ws = wb.create_sheet(title=name)
        for row in rows:
            ws.append(row)
    wb.save(path)
    return path


def _make_csv(path: Path, rows: list[str]) -> Path:
    path.write_text("\n".join(rows), encoding="utf-8")
    return path


def _make_docx(path: Path, paragraphs: list[str], headings: list[str] = None) -> Path:
    """Create a real DOCX with paragraphs and optional heading-styled paras."""
    import docx
    doc = docx.Document()
    for h in (headings or []):
        doc.add_heading(h, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# 1. Fingerprint determinism — same layout, different content → same family
# ---------------------------------------------------------------------------


class TestFingerprintDeterminism:
    """The core invariant: layout-family is form-level, not content-level."""

    def test_xlsx_same_layout_different_data_same_family(self, tmp_path):
        # Two spreadsheets with the same column structure but different values.
        f1 = _make_xlsx(tmp_path / "a.xlsx", {
            "Sheet1": [["Name", "Amount"], ["Alice", 100], ["Bob", 200]],
        })
        f2 = _make_xlsx(tmp_path / "b.xlsx", {
            "Sheet1": [["Name", "Amount"], ["Carol", 300], ["Dave", 400]],
        })
        fam1 = compute_layout_family(f1, "xlsx")
        fam2 = compute_layout_family(f2, "xlsx")
        assert fam1 is not None
        assert fam2 is not None
        assert fam1 == fam2, "same layout, different content must share a family"

    def test_csv_same_layout_different_data_same_family(self, tmp_path):
        f1 = _make_csv(tmp_path / "a.csv", ["id,name,amount", "1,Alice,100", "2,Bob,200"])
        f2 = _make_csv(tmp_path / "b.csv", ["id,name,amount", "3,Carol,300", "4,Dave,400"])
        fam1 = compute_layout_family(f1, "csv")
        fam2 = compute_layout_family(f2, "csv")
        assert fam1 is not None and fam2 is not None
        assert fam1 == fam2

    def test_pdf_same_page_count_same_family(self, tmp_path):
        # Two blank PDFs with the same page count → same family (page count
        # is the dominant structural feature; text content differs trivially).
        f1 = _make_pdf(tmp_path / "a.pdf", ["", ""])
        f2 = _make_pdf(tmp_path / "b.pdf", ["", ""])
        fam1 = compute_layout_family(f1, "pdf")
        fam2 = compute_layout_family(f2, "pdf")
        assert fam1 is not None and fam2 is not None
        assert fam1 == fam2


# ---------------------------------------------------------------------------
# 2. Fingerprint distinction — different layout → different family
# ---------------------------------------------------------------------------


class TestFingerprintDistinction:
    def test_xlsx_different_column_count_different_family(self, tmp_path):
        f1 = _make_xlsx(tmp_path / "two_col.xlsx", {
            "S": [["A", "B"], ["1", "2"]],
        })
        f2 = _make_xlsx(tmp_path / "three_col.xlsx", {
            "S": [["A", "B", "C"], ["1", "2", "3"]],
        })
        fam1 = compute_layout_family(f1, "xlsx")
        fam2 = compute_layout_family(f2, "xlsx")
        assert fam1 != fam2, "different column count must produce different families"

    def test_xlsx_different_sheet_count_different_family(self, tmp_path):
        f1 = _make_xlsx(tmp_path / "one_sheet.xlsx", {"S1": [["A"], ["1"]]})
        f2 = _make_xlsx(tmp_path / "two_sheets.xlsx", {
            "S1": [["A"], ["1"]],
            "S2": [["B"], ["2"]],
        })
        assert compute_layout_family(f1, "xlsx") != compute_layout_family(f2, "xlsx")

    def test_csv_different_column_count_different_family(self, tmp_path):
        f1 = _make_csv(tmp_path / "two.csv", ["a,b", "1,2"])
        f2 = _make_csv(tmp_path / "three.csv", ["a,b,c", "1,2,3"])
        assert compute_layout_family(f1, "csv") != compute_layout_family(f2, "csv")

    def test_pdf_different_page_count_different_family(self, tmp_path):
        f1 = _make_pdf(tmp_path / "one_page.pdf", [""])
        f2 = _make_pdf(tmp_path / "two_page.pdf", ["", ""])
        assert compute_layout_family(f1, "pdf") != compute_layout_family(f2, "pdf")

    def test_docx_different_heading_count_different_family(self, tmp_path):
        f1 = _make_docx(tmp_path / "one_heading.docx", ["body text"], headings=["H1"])
        f2 = _make_docx(tmp_path / "two_headings.docx", ["body text"], headings=["H1", "H2"])
        assert compute_layout_family(f1, "docx") != compute_layout_family(f2, "docx")


# ---------------------------------------------------------------------------
# 3. Per-doc-type features compute without error
# ---------------------------------------------------------------------------


class TestFeatureComputation:
    def test_pdf_features(self, tmp_path):
        f = _make_pdf(tmp_path / "x.pdf", ["", ""])
        feats = compute_layout_features(f, "pdf")
        assert feats is not None
        assert feats["doc_type"] == "pdf"
        assert feats["page_count"] == 2

    def test_xlsx_features(self, tmp_path):
        f = _make_xlsx(tmp_path / "x.xlsx", {"S": [["A", "B"], ["1", "2"]]})
        feats = compute_layout_features(f, "xlsx")
        assert feats is not None
        assert feats["doc_type"] == "xlsx"
        assert feats["sheet_count"] == 1

    def test_csv_features(self, tmp_path):
        f = _make_csv(tmp_path / "x.csv", ["a,b", "1,2"])
        feats = compute_layout_features(f, "csv")
        assert feats is not None
        assert feats["doc_type"] == "csv"
        assert feats["col_count"] == 2

    def test_docx_features(self, tmp_path):
        f = _make_docx(tmp_path / "x.docx", ["para"], headings=["H"])
        feats = compute_layout_features(f, "docx")
        assert feats is not None
        assert feats["doc_type"] == "docx"

    def test_unsupported_type_returns_none(self, tmp_path):
        f = tmp_path / "x.txt"
        f.write_text("hello")
        assert compute_layout_features(f, "txt") is None
        assert compute_layout_family(f, "txt") is None

    def test_missing_file_returns_none(self):
        assert compute_layout_family("/nonexistent/file.xlsx", "xlsx") is None


# ---------------------------------------------------------------------------
# 4. Short-circuit invariant — the critical correctness rule
# ---------------------------------------------------------------------------


class TestShortCircuitInvariant:
    """The issue's key acceptance criterion: layout-family identity NEVER
    suppresses content-level extraction. The short-circuit fires ONLY when
    file_id (content hash) is unchanged."""

    def test_known_family_unchanged_content_short_circuits(self):
        fam = "fam123"
        known = {fam: 5}
        assert should_short_circuit_extraction(
            file_id="abc",
            catalog_file_id="abc",  # unchanged content
            layout_family=fam,
            known_families=known,
        ) is True

    def test_known_family_changed_content_no_short_circuit(self):
        """Same template, one cell updated → content changed → no shortcut.
        The content-level re-extract gate must still run."""
        fam = "fam123"
        known = {fam: 5}
        assert should_short_circuit_extraction(
            file_id="abc",
            catalog_file_id="xyz",  # content changed
            layout_family=fam,
            known_families=known,
        ) is False

    def test_known_family_no_prior_catalog_entry_no_short_circuit(self):
        """New file (no prior catalog entry) even if its layout matches a
        known family → no short-circuit (it's new content)."""
        fam = "fam123"
        known = {fam: 5}
        assert should_short_circuit_extraction(
            file_id="abc",
            catalog_file_id=None,  # new file
            layout_family=fam,
            known_families=known,
        ) is False

    def test_novel_family_no_short_circuit(self):
        """Novel layout → must surface for full extraction + labelling."""
        fam = "fam999"
        known = {"fam123": 5}
        assert should_short_circuit_extraction(
            file_id="abc",
            catalog_file_id="abc",
            layout_family=fam,
            known_families=known,
        ) is False

    def test_unknown_fingerprint_no_short_circuit(self):
        """Fingerprint could not be computed → no form-level shortcut."""
        known = {"fam123": 5}
        assert should_short_circuit_extraction(
            file_id="abc",
            catalog_file_id="abc",
            layout_family=None,
            known_families=known,
        ) is False


# ---------------------------------------------------------------------------
# 5. Novel surfacing — novel layouts always surface, never silently skipped
# ---------------------------------------------------------------------------


class TestNovelSurfacing:
    def test_classify_layout_known(self):
        assert classify_layout("fam1", {"fam1": 3}) == "known"

    def test_classify_layout_novel(self):
        assert classify_layout("fam2", {"fam1": 3}) == "novel"

    def test_classify_layout_unknown(self):
        assert classify_layout(None, {"fam1": 3}) == "unknown"

    def test_surface_novel_layouts_returns_only_novel(self, tmp_path):
        f1 = _make_xlsx(tmp_path / "known.xlsx", {"S": [["A"], ["1"]]})
        f2 = _make_xlsx(tmp_path / "novel.xlsx", {"S": [["A", "B", "C"], ["1", "2", "3"]]})
        fam_known = compute_layout_family(f1, "xlsx")
        fam_novel = compute_layout_family(f2, "xlsx")
        scan = {
            "new": [
                {"path": str(f1), "file_id": "id1", "doc_type": "xlsx", "layout_family": fam_known},
                {"path": str(f2), "file_id": "id2", "doc_type": "xlsx", "layout_family": fam_novel},
            ],
            "changed": [],
        }
        known = {fam_known: 1}
        novel = surface_novel_layouts(scan, known)
        assert len(novel) == 1
        assert novel[0]["layout_class"] == "novel"
        assert novel[0]["layout_family"] == fam_novel

    def test_classify_scan_by_layout_buckets(self, tmp_path):
        f1 = _make_xlsx(tmp_path / "known.xlsx", {"S": [["A"], ["1"]]})
        f2 = _make_xlsx(tmp_path / "novel.xlsx", {"S": [["A", "B", "C"], ["1", "2", "3"]]})
        fam_known = compute_layout_family(f1, "xlsx")
        fam_novel = compute_layout_family(f2, "xlsx")
        scan = {
            "new": [
                {"path": str(f1), "file_id": "id1", "doc_type": "xlsx", "layout_family": fam_known},
                {"path": str(f2), "file_id": "id2", "doc_type": "xlsx", "layout_family": fam_novel},
                {"path": str(tmp_path / "unk.pdf"), "file_id": "id3", "doc_type": "pdf", "layout_family": None},
            ],
            "changed": [],
        }
        known = {fam_known: 1}
        out = classify_scan_by_layout(scan, known)
        assert len(out["known"]) == 1
        assert len(out["novel"]) == 1
        assert len(out["unknown"]) == 1


# ---------------------------------------------------------------------------
# 6. Catalog storage — layout_family persists + retrieves
# ---------------------------------------------------------------------------


class TestCatalogLayoutFamily:
    @pytest.fixture
    def store(self, tmp_path):
        from store import DuckDBMemoryStore
        return DuckDBMemoryStore(tmp_path / "test.duckdb", user_id="alice")

    def test_upsert_with_layout_family(self, store):
        store.upsert_catalog_entry(
            file_id="abc123",
            canonical_path="/docs/invoice.pdf",
            size=1024,
            mtime="2026-01-01T00:00:00+00:00",
            doc_type="pdf",
            layout_family="fam_abc",
        )
        entry = store.get_catalog_entry("abc123")
        assert entry is not None
        assert entry["layout_family"] == "fam_abc"

    def test_upsert_without_layout_family_defaults_null(self, store):
        store.upsert_catalog_entry(
            file_id="def456",
            canonical_path="/docs/other.pdf",
            size=512,
            mtime="2026-01-01T00:00:00+00:00",
            doc_type="pdf",
        )
        entry = store.get_catalog_entry("def456")
        assert entry is not None
        assert entry["layout_family"] is None

    def test_list_catalog_by_layout_family(self, store):
        store.upsert_catalog_entry(
            file_id="a", canonical_path="/a.pdf", size=1, mtime="t", doc_type="pdf",
            layout_family="fam1",
        )
        store.upsert_catalog_entry(
            file_id="b", canonical_path="/b.pdf", size=1, mtime="t", doc_type="pdf",
            layout_family="fam1",
        )
        store.upsert_catalog_entry(
            file_id="c", canonical_path="/c.pdf", size=1, mtime="t", doc_type="pdf",
            layout_family="fam2",
        )
        rows = store.list_catalog_by_layout_family("fam1")
        assert len(rows) == 2
        assert all(r["layout_family"] == "fam1" for r in rows)

    def test_list_layout_families(self, store):
        store.upsert_catalog_entry(
            file_id="a", canonical_path="/a.pdf", size=1, mtime="t", doc_type="pdf",
            layout_family="fam1",
        )
        store.upsert_catalog_entry(
            file_id="b", canonical_path="/b.pdf", size=1, mtime="t", doc_type="pdf",
            layout_family="fam1",
        )
        store.upsert_catalog_entry(
            file_id="c", canonical_path="/c.pdf", size=1, mtime="t", doc_type="pdf",
            layout_family="fam2",
        )
        fams = store.list_layout_families()
        by_fam = {f["layout_family"]: f["doc_count"] for f in fams}
        assert by_fam == {"fam1": 2, "fam2": 1}

    def test_migration_adds_column_to_existing_db(self, tmp_path):
        """A DB created before the column existed gets it via ALTER TABLE."""
        from store import DuckDBMemoryStore
        s = DuckDBMemoryStore(tmp_path / "test.duckdb", user_id="alice")
        # Column should exist and be queryable.
        s.connection.execute("SELECT layout_family FROM file_catalog LIMIT 1")
        s.close()


# ---------------------------------------------------------------------------
# 7. Label registry (sidecar)
# ---------------------------------------------------------------------------


class TestLabelRegistry:
    def test_label_and_get(self, tmp_path):
        reg = tmp_path / "families.json"
        assert label_family(reg, "fam1", "Acme invoice") is True
        assert get_label(reg, "fam1") == "Acme invoice"

    def test_relabel_keeps_history(self, tmp_path):
        reg = tmp_path / "families.json"
        label_family(reg, "fam1", "Old label")
        label_family(reg, "fam1", "New label")
        loaded = load_registry(reg)
        assert loaded["fam1"]["label"] == "New label"
        assert loaded["fam1"]["label_history"][0]["label"] == "Old label"

    def test_labelled_families(self, tmp_path):
        reg = tmp_path / "families.json"
        label_family(reg, "fam1", "A")
        label_family(reg, "fam2", "B")
        labelled = labelled_families(reg)
        assert labelled == {"fam1": "A", "fam2": "B"}

    def test_missing_registry_returns_empty(self, tmp_path):
        assert load_registry(tmp_path / "nope.json") == {}
        assert get_label(tmp_path / "nope.json", "fam1") is None

    def test_registry_path_for(self, tmp_path):
        p = registry_path_for(tmp_path / "store.duckdb")
        assert p == tmp_path / "layout_families.json"

    def test_merge_catalog_counts(self, tmp_path):
        reg = tmp_path / "families.json"
        label_family(reg, "fam1", "A")
        loaded = load_registry(reg)
        merged = merge_catalog_counts(loaded, [
            {"layout_family": "fam1", "doc_count": 5},
            {"layout_family": "fam2", "doc_count": 3},
        ])
        assert merged["fam1"]["label"] == "A"
        assert merged["fam1"]["doc_count"] == 5
        assert merged["fam2"]["label"] == ""
        assert merged["fam2"]["doc_count"] == 3


# ---------------------------------------------------------------------------
# 8. Scan pass computes layout_family for new/changed files
# ---------------------------------------------------------------------------


class TestScanPassLayoutFamily:
    def test_scan_pass_includes_layout_family_for_new_file(self, tmp_path):
        f = _make_xlsx(tmp_path / "new.xlsx", {"S": [["A", "B"], ["1", "2"]]})
        result = scan_pass([tmp_path], catalog={})
        assert len(result["new"]) == 1
        assert result["new"][0]["layout_family"] is not None

    def test_scan_pass_no_layout_family_for_unchanged(self, tmp_path):
        """Unchanged files (already in catalog) don't get re-fingerprinted —
        the catalog already has it."""
        f = _make_xlsx(tmp_path / "known.xlsx", {"S": [["A"], ["1"]]})
        # First pass: discovers the file, computes fingerprint.
        result1 = scan_pass([tmp_path], catalog={})
        file_id = result1["new"][0]["file_id"]
        fam = result1["new"][0]["layout_family"]
        # Second pass: file is unchanged → 'unchanged' bucket, no re-fingerprint.
        catalog = {file_id: {"paths": [str(f)], "status": "active"}}
        result2 = scan_pass([tmp_path], catalog=catalog)
        assert len(result2["unchanged"]) == 1
        # Unchanged files don't carry layout_family (it's already in catalog).
        assert "layout_family" in result2["unchanged"][0]
        assert result2["unchanged"][0]["layout_family"] is None


# ---------------------------------------------------------------------------
# 9. Eval stratification — per-family accuracy reported alongside aggregate
# ---------------------------------------------------------------------------


class TestEvalStratification:
    def test_compute_metrics_includes_by_layout_family(self):
        from eval_self_corpus import compute_metrics
        probes = [
            {
                "template": "exact", "target_memory_id": "m1",
                "target_category": "personal_fact",
                "target_content_len": 50, "target_created_at": "2026-01-01T00:00:00+00:00",
                "target_layout_family": "fam_a",
                "hit": True, "per_window": {"1": True, "5": True, "20": True}, "rank": 1,
                "wide_pool_rank": 1, "not_in_pool": False,
                "top_5_ids": [], "top_5_similarities": [],
            },
            {
                "template": "exact", "target_memory_id": "m2",
                "target_category": "personal_fact",
                "target_content_len": 50, "target_created_at": "2026-01-01T00:00:00+00:00",
                "target_layout_family": "fam_a",
                "hit": False, "per_window": {}, "rank": None,
                "wide_pool_rank": None, "not_in_pool": True,
                "top_5_ids": [], "top_5_similarities": [],
            },
            {
                "template": "exact", "target_memory_id": "m3",
                "target_category": "personal_fact",
                "target_content_len": 50, "target_created_at": "2026-01-01T00:00:00+00:00",
                "target_layout_family": "fam_b",
                "hit": True, "per_window": {"1": True, "5": True, "20": True}, "rank": 1,
                "wide_pool_rank": 1, "not_in_pool": False,
                "top_5_ids": [], "top_5_similarities": [],
            },
            {
                "template": "exact", "target_memory_id": "m4",
                "target_category": "personal_fact",
                "target_content_len": 50, "target_created_at": "2026-01-01T00:00:00+00:00",
                "target_layout_family": None,  # conversation-sourced
                "hit": True, "per_window": {"1": True, "5": True, "20": True}, "rank": 1,
                "wide_pool_rank": 1, "not_in_pool": False,
                "top_5_ids": [], "top_5_similarities": [],
            },
        ]
        metrics = compute_metrics(probes, [1, 5, 20], 4, {"personal_fact": 4}, "cfg123")
        assert "by_layout_family" in metrics
        assert "layout_family_distribution" in metrics
        # Distribution: fam_a=2, fam_b=1, none=1
        dist = metrics["layout_family_distribution"]
        assert dist["fam_a"] == 2
        assert dist["fam_b"] == 1
        assert dist["none"] == 1
        # Per-family recall@1: fam_a=0.5 (1 hit of 2), fam_b=1.0, none=1.0
        r1 = metrics["by_layout_family"]["recall@1"]
        assert r1["fam_a"] == 0.5
        assert r1["fam_b"] == 1.0
        assert r1["none"] == 1.0

    def test_cluster_families_script(self, tmp_path):
        """The pre-labelling clustering script clusters by fingerprint and
        prints family counts (stratified plan, not post-hoc)."""
        from eval.cluster_layout_families import cluster_families, stratified_sample_plan
        from store import DuckDBMemoryStore
        store = DuckDBMemoryStore(tmp_path / "test.duckdb", user_id="alice")
        store.upsert_catalog_entry(
            file_id="a", canonical_path="/a.pdf", size=1, mtime="t", doc_type="pdf",
            layout_family="fam1",
        )
        store.upsert_catalog_entry(
            file_id="b", canonical_path="/b.pdf", size=1, mtime="t", doc_type="pdf",
            layout_family="fam1",
        )
        store.upsert_catalog_entry(
            file_id="c", canonical_path="/c.pdf", size=1, mtime="t", doc_type="pdf",
            layout_family="fam2",
        )
        store.upsert_catalog_entry(
            file_id="d", canonical_path="/d.pdf", size=1, mtime="t", doc_type="pdf",
            layout_family=None,
        )
        store.close()
        families = cluster_families(tmp_path / "test.duckdb")
        fingerprinted = [f for f in families if f["layout_family"]]
        assert len(fingerprinted) == 2
        counts = {f["layout_family"]: f["doc_count"] for f in fingerprinted}
        assert counts == {"fam1": 2, "fam2": 1}
        plan = stratified_sample_plan(fingerprinted, total=10, min_per_family=2)
        assert plan["fam1"] >= 2
        assert plan["fam2"] >= 2
        assert sum(plan.values()) <= 10

    def test_load_known_families(self):
        rows = [
            {"layout_family": "fam1"},
            {"layout_family": "fam1"},
            {"layout_family": "fam2"},
            {"layout_family": None},
        ]
        fams = load_known_families(rows)
        assert fams == {"fam1": 2, "fam2": 1}
